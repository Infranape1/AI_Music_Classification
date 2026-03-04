from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
import librosa
import numpy as np
import pickle
import os
from werkzeug.utils import secure_filename
from keras.models import load_model
from reportlab.pdfgen import canvas
from docx import Document

# -----------------------------
# Flask Setup
# -----------------------------
app = Flask(__name__)
CORS(app)

# -----------------------------
# Model & Encoder Load
# -----------------------------
MODEL_FILE = "music_genre_model.h5"       # Make sure your model is saved as .h5
ENCODER_FILE = "label_encoder.pkl"

if not os.path.exists(MODEL_FILE) or not os.path.exists(ENCODER_FILE):
    raise FileNotFoundError("Model or encoder file not found. Upload them in the project root.")

model = load_model(MODEL_FILE, compile=False)
with open(ENCODER_FILE, "rb") as f:
    encoder = pickle.load(f)

# -----------------------------
# Upload Folder & Allowed Extensions
# -----------------------------
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
ALLOWED_EXTENSIONS = {"wav", "mp3", "flac", "ogg", "m4a"}

MODEL_ACCURACY = 0.55
MODEL_F1_SCORE = 0.52

# -----------------------------
# Helper Functions
# -----------------------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_mfcc(file_path, max_pad_len=130):
    audio, sr = librosa.load(file_path, duration=30)
    mfcc = librosa.feature.mfcc(y=audio, sr=sr, n_mfcc=40)
    mfcc = (mfcc - np.mean(mfcc)) / (np.std(mfcc) + 1e-10)
    pad_width = max_pad_len - mfcc.shape[1]
    if pad_width > 0:
        mfcc = np.pad(mfcc, ((0,0),(0,pad_width)), mode='constant')
    else:
        mfcc = mfcc[:, :max_pad_len]
    return mfcc

# -----------------------------
# Serve Frontend
# -----------------------------
@app.route("/")
def home():
    return send_from_directory(".", "index.html")

# -----------------------------
# Prediction API
# -----------------------------
@app.route("/predict", methods=["POST"])
def predict():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Unsupported file format"}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(file_path)

    try:
        mfcc = extract_mfcc(file_path)
        mfcc = mfcc.reshape(1, 40, 130, 1)
        prediction = model.predict(mfcc)
        predicted_index = np.argmax(prediction)
        genre = encoder.inverse_transform([predicted_index])[0]
        confidence = float(np.max(prediction)) * 100

        # Prepare analysis data
        analysis = {
            "audio_file": filename,
            "genre": genre,
            "confidence": round(confidence,2),
            "accuracy": MODEL_ACCURACY,
            "f1_score": MODEL_F1_SCORE
        }

        # Save PDF
        pdf_path = os.path.join(UPLOAD_FOLDER, f"{filename}_analysis.pdf")
        c = canvas.Canvas(pdf_path)
        c.setFont("Helvetica", 14)
        y = 750
        for k,v in analysis.items():
            c.drawString(50, y, f"{k}: {v}")
            y -= 25
        c.save()

        # Save DOCX
        docx_path = os.path.join(UPLOAD_FOLDER, f"{filename}_analysis.docx")
        doc = Document()
        doc.add_heading("Music Genre Analysis", level=1)
        for k,v in analysis.items():
            doc.add_paragraph(f"{k}: {v}")
        doc.save(docx_path)

        os.remove(file_path)

        return jsonify({
            "genre": genre,
            "confidence": round(confidence,2),
            "accuracy": MODEL_ACCURACY,
            "f1_score": MODEL_F1_SCORE,
            "pdf_file": f"{filename}_analysis.pdf",
            "docx_file": f"{filename}_analysis.docx"
        })

    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        return jsonify({"error": str(e)}), 500

# -----------------------------
# Download API
# -----------------------------
@app.route("/download/<file_name>")
def download(file_name):
    return send_file(os.path.join(UPLOAD_FOLDER, file_name), as_attachment=True)

# -----------------------------
# Run Flask App
# -----------------------------
if __name__ == "__main__":
    import os
    # Get the port number from the environment variable, default to 7860 if not set
    port = int(os.environ.get("PORT", 7860))
    # Run the Flask app on all interfaces
    app.run(host="0.0.0.0", port=port, debug=True)